from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

# ── 字体常量 ──────────────────────────────────────────
ZH_FONT   = 'SimSun'       # 宋体
EN_FONT   = 'Times New Roman'
TITLE_PT  = 16             # 3号
H1_PT     = 14             # 四号
H2_PT     = 13             # 小三偏下
H3_PT     = 12             # 小四
BODY_PT   = 12             # 小四
HINT_PT   = 10.5           # 提示文字
INDENT_PT = 24             # 首行2字符缩进（2×12pt）

# ── 颜色 ─────────────────────────────────────────────
CLR_TITLE    = RGBColor(0x1A, 0x1A, 0x2E)   # 深蓝黑
CLR_H1       = RGBColor(0x1A, 0x3A, 0x5C)   # 深海蓝
CLR_H2       = RGBColor(0x1F, 0x5C, 0x8B)   # 中蓝
CLR_H3       = RGBColor(0x2E, 0x2E, 0x2E)   # 近黑
CLR_HINT     = RGBColor(0x88, 0x88, 0x88)   # 灰
CLR_ONELINER = RGBColor(0x1A, 0x3A, 0x5C)   # 与H1同色
CLR_BODY     = RGBColor(0x1A, 0x1A, 0x1A)   # 正文近黑


def _set_run_zh(run, size, bold=False, italic=False, color=None):
    """统一设置中文字体（宋体）"""
    run.font.name = EN_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    EN_FONT)
    rFonts.set(qn('w:hAnsi'),    EN_FONT)
    rFonts.set(qn('w:eastAsia'), ZH_FONT)
    rFonts.set(qn('w:cs'),       ZH_FONT)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def _set_para_spacing(para, line_spacing=1.5,
                       space_before=0, space_after=6,
                       first_indent=None, left_indent=None,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT):
    pf = para.paragraph_format
    pf.alignment = alignment
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if first_indent is not None:
        pf.first_line_indent = Pt(first_indent)
    if left_indent is not None:
        pf.left_indent = Pt(left_indent)


def add_hr(doc, color='C0C0C0', thickness=4):
    """插入分隔线"""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    str(thickness))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    _set_para_spacing(para, space_before=0, space_after=2)


def add_title_block(doc):
    """文档主标题区域"""
    # 留白
    sp = doc.add_paragraph()
    _set_para_spacing(sp, space_before=0, space_after=4)

    # 主标题
    p = doc.add_paragraph()
    _set_para_spacing(p, space_before=0, space_after=6,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.2)
    r = p.add_run('KuberCloud ANI 产品定义')
    _set_run_zh(r, TITLE_PT, bold=True, color=CLR_TITLE)

    # 副标题
    p2 = doc.add_paragraph()
    _set_para_spacing(p2, space_before=0, space_after=4,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.2)
    r2 = p2.add_run('KuberCloud AI-Native Infrastructure  ·  AI专有云')
    _set_run_zh(r2, HINT_PT + 0.5, color=CLR_HINT)

    # 分隔线
    add_hr(doc, color='4A90D9', thickness=6)


def add_h1(doc, text):
    """一级大标题（如：一句话定义 / 凝练版 / 完整版）"""
    p = doc.add_paragraph()
    _set_para_spacing(p, space_before=14, space_after=4, line_spacing=1.3)
    r = p.add_run(text)
    _set_run_zh(r, H1_PT, bold=True, color=CLR_H1)
    # 段落下方细线
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '2')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '4A90D9')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_h2(doc, text):
    """二级标题（如：同源共生，全态承载 / 统一算力底座）"""
    p = doc.add_paragraph()
    _set_para_spacing(p, space_before=10, space_after=3, line_spacing=1.3)
    r = p.add_run(text)
    _set_run_zh(r, H2_PT, bold=True, color=CLR_H2)


def add_h3(doc, text):
    """三级标题（如：· 传统业务应用）"""
    p = doc.add_paragraph()
    _set_para_spacing(p, space_before=8, space_after=2,
                      left_indent=12, line_spacing=1.3)
    r = p.add_run(text)
    _set_run_zh(r, H3_PT, bold=True, color=CLR_H3)


def add_hint(doc, text):
    """适用场景提示行"""
    p = doc.add_paragraph()
    _set_para_spacing(p, space_before=0, space_after=6, line_spacing=1.2)
    r = p.add_run(text)
    _set_run_zh(r, HINT_PT, italic=True, color=CLR_HINT)


def add_body(doc, text, first_indent=True):
    """标准正文段落：1.5倍行距 + 首行缩进2字符"""
    p = doc.add_paragraph()
    _set_para_spacing(p, space_before=0, space_after=5, line_spacing=1.5,
                      first_indent=INDENT_PT if first_indent else None)
    r = p.add_run(text)
    _set_run_zh(r, BODY_PT, color=CLR_BODY)


def add_oneliner(doc, text):
    """一句话定义特殊样式：居中，稍大，带引号框感"""
    p = doc.add_paragraph()
    _set_para_spacing(p, space_before=6, space_after=6, line_spacing=1.5,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(text)
    _set_run_zh(r, H2_PT, bold=True, color=CLR_ONELINER)


def add_bullet(doc, label, content):
    """带加粗标签的要点段落（安全沙箱：...）"""
    p = doc.add_paragraph()
    _set_para_spacing(p, space_before=3, space_after=4, line_spacing=1.5,
                      left_indent=24, first_indent=-12)
    # 圆点
    dot = p.add_run('● ')
    _set_run_zh(dot, BODY_PT, bold=True, color=CLR_H2)
    # 加粗标签
    r_label = p.add_run(label)
    _set_run_zh(r_label, BODY_PT, bold=True, color=CLR_BODY)
    # 正文内容
    r_body = p.add_run(content)
    _set_run_zh(r_body, BODY_PT, color=CLR_BODY)


def add_meta(doc, key, value):
    """封面元信息行"""
    p = doc.add_paragraph()
    _set_para_spacing(p, space_before=0, space_after=2, line_spacing=1.3)
    rk = p.add_run(key)
    _set_run_zh(rk, HINT_PT, bold=True, color=CLR_HINT)
    rv = p.add_run(value)
    _set_run_zh(rv, HINT_PT, color=CLR_HINT)


# ═══════════════════════════════════════════════════════
#  正式构建文档
# ═══════════════════════════════════════════════════════
doc = Document()

# ── 页面设置 ──────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.8)
section.bottom_margin = Cm(2.8)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(3.0)

# ── 标题区 ────────────────────────────────────────────
add_title_block(doc)

# 元信息
sp = doc.add_paragraph()
_set_para_spacing(sp, space_before=6, space_after=2)
add_meta(doc, '产品全称：', 'KuberCloud AI-Native Infrastructure')
add_meta(doc, '中文名称：', 'AI专有云')
add_meta(doc, '品牌简写：', 'KuberCloud ANI')
add_meta(doc, '文档版本：', f'V7  |  {datetime.now().strftime("%Y年%m月")}')
add_hr(doc)

# ── 一句话定义 ────────────────────────────────────────
add_h1(doc, '一句话定义')
add_oneliner(doc, 'AI专有云，是为企业级 AI 应用而构建的专属智能算力云平台。')
add_hr(doc, color='E0E0E0', thickness=2)

# ── 凝练版 ────────────────────────────────────────────
add_h1(doc, '凝练版')
add_hint(doc, '适用场景：官网首屏 / 产品手册 / Pitch Deck')
add_body(doc,
    'KuberCloud ANI（全称：KuberCloud AI-Native Infrastructure，AI专有云）是常青云面向 AI 原生时代构建的企业级专属算力云平台，'
    '以独立资源池、物理隔离与全栈专属的形式交付。以"同源共生，全态承载"为核心架构哲学，在统一算力底座之上，并行支撑传统业务应用、'
    'AI 增强型应用与 AI 原生应用三种工作负载形态，同时提供覆盖模型全生命周期的 MLOps 工程环境。使企业无论处于哪种技术演进阶段，'
    '均可在安全、可控、合规的专属环境中规模化构建和运行业务应用。')
add_hr(doc, color='E0E0E0', thickness=2)

# ── 完整版 ────────────────────────────────────────────
add_h1(doc, '完整版')
add_hint(doc, '适用场景：白皮书 / 招标文件 / 战略合作协议')
add_body(doc,
    'KuberCloud ANI（全称：KuberCloud AI-Native Infrastructure，AI专有云）是常青云面向 AI 原生时代构建的企业级智能算力云平台，'
    '以独立资源池、物理隔离与全栈专属的方式交付，将云计算、异构算力调度、模型工程平台、AI 原生应用运行时与安全治理能力统一集成，'
    '部署于客户指定的专属资源环境。在保障数据主权、安全隔离与监管合规的前提下，兼具公共云级别的弹性编排与敏捷运营效率。')

# 同源共生
add_h2(doc, '同源共生，全态承载')
add_body(doc,
    'KuberCloud ANI 以"同源共生，全态承载"为核心架构哲学：所有工作负载共享同一算力底座（同源），传统业务应用、AI 增强型应用与 '
    'AI 原生应用在同一平台上并存演进（共生），平台无边界地承载企业当下与未来的全部应用形态（全态承载）。')
add_body(doc,
    '企业的技术演进从不是一步到位的跃迁。按照 2/8 原则，80% 的业务仍需在稳定的现有架构上持续运行，AI 能力的植入是渐进式的。'
    'KuberCloud ANI 不强迫客户做非此即彼的技术选择，而是让三种应用架构形态在同一平台上共存共生、按需演进。')

# 统一算力底座
add_h2(doc, '统一算力底座')
add_body(doc,
    'KuberCloud ANI 对 GPU、NPU、FPGA 等 AI 加速硬件，以及 CPU、虚拟机（VM）、裸金属服务器等通用计算资源进行统一抽象、'
    '声明式调度与弹性编排，支持跨厂商、跨代际的算力资源动态分配与高效复用。无论传统业务还是 AI 原生工作负载，均从同一算力底座'
    '按需取用，如同水电基础设施般随取随用，充分释放存量硬件资产的投资价值。')

# 三态共生
add_h2(doc, '三态共生')

add_h3(doc, '· 传统业务应用')
add_body(doc,
    '支持虚拟机（VM）与裸金属服务器的标准化交付与全生命周期管理，为现有业务系统提供稳定、可靠的运行环境，保障企业核心业务在向 AI '
    '演进过程中的连续性与稳定性。同时提供容器化迁移路径，帮助传统应用以最小改造成本逐步向云原生架构过渡。')

add_h3(doc, '· AI 增强型应用')
add_body(doc,
    '面向正在将 AI 能力融入现有业务系统的企业，提供传统应用架构与 AI 服务的无缝集成环境。企业无需重写现有系统，即可通过模型推理'
    '服务集成、智能组件嵌入等方式，在现有业务流程中引入 AI 决策、智能推荐、自动化处理等能力，实现渐进式 AI 化。')

add_h3(doc, '· AI 原生应用')
add_body(doc,
    '为 AI Agent、多智能体系统（Multi-Agent Systems）与全新 AI 原生应用提供完整的开发、运行与治理能力。')

add_bullet(doc, '安全沙箱执行环境：',
    '提供进程级、容器级与虚拟机级的多层安全沙箱执行环境，为 AI Agent 的代码执行、工具调用与外部交互构建严格的安全隔离边界，'
    '有效防范提示注入（Prompt Injection）等 AI 特有安全威胁，满足企业级安全合规要求。沙箱隔离级别可依据工作负载类型与租户安全'
    '策略灵活配置，在安全隔离与运行性能之间实现最优平衡。')

add_bullet(doc, 'AI Agent 编排与治理：',
    '支持 AI Agent 工作流的动态调度与全生命周期管理，提供细粒度的资源配额、网络隔离策略与访问权限控制，确保多租户环境下不同 '
    'Agent 工作负载间的安全边界与资源公平性。')

add_bullet(doc, 'AI 原生应用开发与运行：',
    '为 AI Agent 框架及行业 AI 原生应用提供标准化运行环境、统一 API 入口管理、全链路可观测性（追踪 / 日志 / 指标）与运维管理'
    '能力，使 AI 原生应用具备生产级的稳定性、可扩展性与可运营性。')

# 模型工程平台
add_h2(doc, '模型工程平台')
add_body(doc,
    '提供覆盖数据治理、模型训练、微调、评测、推理部署与运维管理的一站式 MLOps 工程环境，为 AI 增强型应用与 AI 原生应用提供模型'
    '开发与服务能力支撑，帮助客户完成从算力建设到模型落地、从实验验证到生产规模化运行的完整链路，消除跨平台集成的工程负担与重复'
    '建设成本。')

# 范式重构
add_h2(doc, '范式重构：从"AI on Cloud"到"Cloud for AI"')
add_body(doc,
    'KuberCloud ANI 不只是将 AI 工作负载搬上云端，而是以 AI 的训练、推理、Agent 调度、沙箱安全与行业场景需求为核心，重新组织'
    '云的算力、平台与服务能力——这是 AI 基础设施底层逻辑的根本性重构。通过专属化部署、平台化能力与场景化服务，KuberCloud ANI '
    '大幅降低企业 AI 规模化落地的门槛，使行业客户聚焦业务创新，实现智能算力从资源投入到业务价值的高效转化与完整闭环。')

# ── 保存 ──────────────────────────────────────────────
output_path = '/Users/zhangfan/ANI/KuberCloud-ANI-产品定义.docx'
doc.save(output_path)
print(f'Word 文档已生成：{output_path}')
